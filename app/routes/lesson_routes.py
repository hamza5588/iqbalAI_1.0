from flask import Blueprint, request, jsonify, session, send_file, after_this_request, render_template
from app.models.models import UserModel, LessonModel
from app.models.database_models import Lesson as DBLesson
from app.services.lesson_service import LessonService
from app.utils.decorators import login_required, teacher_required, student_required
from app.utils.db import get_db
from werkzeug.datastructures import FileStorage
import logging
import os
from io import BytesIO
import tempfile

logger = logging.getLogger(__name__)

# Create a custom logger for lesson checks
lesson_check_logger = logging.getLogger('lesson_check')
lesson_check_logger.setLevel(logging.DEBUG)

# Ensure logs directory exists
os.makedirs('logs', exist_ok=True)

lesson_check_handler = logging.FileHandler('logs/lesson_check.log')
lesson_check_handler.setLevel(logging.DEBUG)
lesson_check_formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
lesson_check_handler.setFormatter(lesson_check_formatter)
lesson_check_logger.addHandler(lesson_check_handler)
lesson_check_logger.info("=" * 80)
lesson_check_logger.info("Lesson Check Logger Initialized")
lesson_check_logger.info("=" * 80)

bp = Blueprint('lesson_routes', __name__)

@bp.route('/create_lesson', methods=['POST'])
# @teacher_required
def create_lesson():
    """Create a new lesson or answer a question based on user input"""
    try:
        # Get form data for lesson configuration
        lesson_title = request.form.get('lessonTitle', '')
        lesson_prompt = request.form.get('lessonPrompt', '')
        focus_area = request.form.get('focusArea', '')
        grade_level = request.form.get('gradeLevel', '')
        additional_notes = request.form.get('additionalNotes', '')
        
        # Get API key from session
        api_key = session.get('groq_api_key')
        if not api_key:
            return jsonify({'error': 'API key not configured. Please set your API key first.'}), 400
        
        # Initialize lesson service
        lesson_service = LessonService(api_key=api_key)
        
        # Check if file is provided - if yes, always treat as lesson generation
        if 'file' in request.files and request.files['file'].filename != '':
            # User uploaded a file - proceed with lesson generation regardless of prompt type
            logger.info("File uploaded - proceeding with lesson generation")
            file = request.files['file']
            
            # Validate required fields for lesson generation (prompt is now optional since we use interactive chat)
            if not lesson_title or not focus_area or not grade_level:
                return jsonify({'error': 'All required fields must be filled for lesson generation'}), 400
            
            # Check if lesson title already exists for this teacher
            if LessonModel.check_title_exists(session['user_id'], lesson_title):
                return jsonify({'error': 'This lesson title is already used. Please choose a different title.'}), 400
            
            # Check file type
            allowed_extensions = {'.pdf', '.doc', '.docx', '.txt'}
            file_ext = os.path.splitext(file.filename.lower())[1]
            if file_ext not in allowed_extensions:
                return jsonify({'error': 'File type not supported. Please upload PDF, DOC, DOCX, or TXT files.'}), 400
            
            # Get extraction toggle values
            table_extraction = request.form.get('tableExtraction', 'false').lower() == 'true'
            image_extraction = request.form.get('imageExtraction', 'false').lower() == 'true'
            
            # Prepare lesson details
            lesson_details = {
                'lesson_title': lesson_title,
                'lesson_prompt': lesson_prompt,
                'focus_area': focus_area,
                'grade_level': grade_level,
                'additional_notes': additional_notes,
                'table_extraction': table_extraction,
                'image_extraction': image_extraction
            }
            
            # Process the file first to create vector store
            process_result = lesson_service.process_file(file, lesson_details)
            
            if 'error' in process_result:
                return jsonify({'error': process_result['error']}), 400
            
            # Create a draft lesson entry in database
            # Use a default summary since prompt is no longer required
            summary_text = f"Lesson on {focus_area} for {grade_level} grade students"
            if additional_notes:
                summary_text = additional_notes[:200] + "..." if len(additional_notes) > 200 else additional_notes
            
            lesson_id = LessonModel.create_lesson(
                teacher_id=session['user_id'],
                title=lesson_title,
                summary=summary_text,
                learning_objectives='',
                content='',  # Will be filled when lesson is complete
                grade_level=grade_level,
                focus_area=focus_area,
                is_public=True  # Make lessons public by default
            )
            
            if not lesson_id:
                return jsonify({'error': 'Failed to save lesson to database'}), 500
            
            # Get the greeting message from process_result (no LLM call was made)
            greeting_message = process_result.get('lesson', 'Your file has been uploaded successfully.')
            
            # Get the lesson response
            lesson_response = LessonModel.get_lesson_by_id(lesson_id)
            lesson_response['title'] = lesson_title
            lesson_response['grade_level'] = grade_level
            lesson_response['focus_area'] = focus_area
            lesson_response['isFinalized'] = False
            
            # Return greeting message without calling LLM
            # LLM will only be called when user sends a query via interactive_chat endpoint
            return jsonify({
                'success': True,
                'response_type': 'file_uploaded',
                'lesson_id': lesson_id,
                'lesson': lesson_response,
                'ai_response': greeting_message,
                'complete_lesson': 'no',
                'message': 'File uploaded successfully! You can now start chatting to create your lesson.',
                'file_processed': True,
                'filename': process_result.get('filename_processed', file.filename)
            })
        
        else:
            # No file uploaded - use additional_notes or lesson_prompt if provided, otherwise require file
            query_text = additional_notes or lesson_prompt or ''
            
            if query_text:
                # Analyze query to determine if it's a question or lesson generation request
                query_analysis = lesson_service.analyze_user_query(query_text)
                logger.info(f"Query analysis: {query_analysis}")
                
                if query_analysis['query_type'] == 'QUESTION':
                    # User is asking a question - answer it using available lessons
                    logger.info("User query detected as question - answering using available lessons")
                    
                    # Get available lesson IDs for context
                    available_lessons = lesson_service._get_available_lesson_ids()
                    
                    # Answer the question using available lesson content
                    answer_result = lesson_service.answer_general_question(query_text, available_lessons)
                    
                    return jsonify({
                        'success': True,
                        'response_type': 'question_answer',
                        'answer': answer_result['answer'],
                        'source': answer_result.get('source', 'unknown'),
                        'confidence': answer_result.get('confidence', 0.0),
                        'lesson_context': answer_result.get('lesson_context', ''),
                        'relevant_lessons': answer_result.get('relevant_lessons', []),
                        'query_analysis': query_analysis,
                        'message': 'Question answered based on available lesson content'
                    })
            
            # User wants to generate a lesson but no file provided
            return jsonify({'error': 'File is required for lesson generation. Please upload a PDF, DOC, DOCX, or TXT file.'}), 400
        
    except Exception as e:
        logger.error(f"Lesson creation/query error: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to process request: {str(e)}'}), 500

@bp.route('/ask_question_general', methods=['POST'])
@login_required
def ask_general_question():
    """Ask a general question that will be answered using available lesson content"""
    try:
        data = request.get_json()
        question = data.get('question', '').strip()
        
        if not question:
            return jsonify({'error': 'Question is required'}), 400
        
        # Get API key from session
        api_key = session.get('groq_api_key')
        if not api_key:
            return jsonify({'error': 'API key not configured. Please set your API key first.'}), 400
        
        # Initialize lesson service
        lesson_service = LessonService(api_key=api_key)
        
        # Get conversation history from all lessons for context
        from app.models.models import LessonChatHistory
        user_id = session['user_id']
        
        # Get recent conversation history from all lessons
        conversation_history = []
        available_lessons = lesson_service._get_available_lesson_ids()
        
        # Get recent Q&A from all lessons (last 5 from each lesson)
        for lesson_id in available_lessons[:5]:  # Limit to first 5 lessons to avoid too much context
            try:
                lesson_history = LessonChatHistory.get_lesson_chat_history(lesson_id, user_id)
                conversation_history.extend(lesson_history[-2:])  # Last 2 from each lesson
            except Exception as e:
                logger.warning(f"Error getting history for lesson {lesson_id}: {str(e)}")
                continue
        
        # Sort by timestamp and take most recent
        conversation_history = sorted(conversation_history, key=lambda x: x.get('created_at', ''), reverse=True)[:10]
        
        # Analyze the user's query to determine intent
        query_analysis = lesson_service.analyze_user_query(question)
        logger.info(f"General query analysis: {query_analysis}")
        
        if not available_lessons:
            return jsonify({
                'success': True,
                'response_type': 'no_content',
                'answer': 'I don\'t have any lesson content available to answer your question. Please upload some educational materials first.',
                'source': 'no_content',
                'confidence': 0.0,
                'query_analysis': query_analysis,
                'message': 'No lesson content available'
            })
        
        # Answer the question using available lesson content with conversation context
        answer_result = lesson_service.answer_general_question(question, available_lessons, conversation_history)
        
        return jsonify({
            'success': True,
            'response_type': 'question_answer',
            'answer': answer_result['answer'],
            'source': answer_result.get('source', 'unknown'),
            'confidence': answer_result.get('confidence', 0.0),
            'lesson_context': answer_result.get('lesson_context', ''),
            'relevant_lessons': answer_result.get('relevant_lessons', []),
            'query_analysis': query_analysis,
            'message': 'Question answered based on available lesson content'
        })
        
    except Exception as e:
        logger.error(f"Error answering general question: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to answer question: {str(e)}'}), 500

@bp.route('/my_lessons', methods=['GET'])
@teacher_required
def get_my_lessons():
    """Get all lessons created by the current teacher"""
    try:
        lessons = LessonModel.get_lessons_by_teacher(session['user_id'])
        return jsonify({
            'success': True,
            'lessons': lessons
        })
    except Exception as e:
        logger.error(f"Error getting teacher lessons: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to get lessons: {str(e)}'}), 500

@bp.route('/browse_lessons', methods=['GET'])
@student_required
def browse_lessons():
    """Browse available lessons for students"""
    try:
        grade_level = request.args.get('grade_level')
        focus_area = request.args.get('focus_area')
        
        lessons = LessonModel.get_public_latest_lessons(grade_level=grade_level, focus_area=focus_area)
        return jsonify({
            'success': True,
            'lessons': lessons
        })
    except Exception as e:
        logger.error(f"Error browsing lessons: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to browse lessons: {str(e)}'}), 500

@bp.route('/search_lessons', methods=['GET'])
@login_required
def search_lessons():
    """Search lessons for both teachers and students"""
    try:
        search_term = request.args.get('q', '')
        grade_level = request.args.get('grade_level')
        
        if not search_term:
            return jsonify({'error': 'Search term is required'}), 400
        
        lessons = LessonModel.search_lessons(search_term, grade_level=grade_level)
        return jsonify({
            'success': True,
            'lessons': lessons
        })
    except Exception as e:
        logger.error(f"Error searching lessons: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to search lessons: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>', methods=['GET'])
@login_required
def get_lesson(lesson_id):
    """Get a specific lesson by ID"""
    try:
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        # Check if user can access this lesson
        user_id = session.get('user_id')
        user_role = session.get('role', 'student')
        lesson_teacher_id = lesson.get('teacher_id')
        is_public = lesson.get('is_public', False)
        
        # Users can access their own lessons (regardless of role)
        if lesson_teacher_id == user_id:
            pass  # Allow access - user owns this lesson
        # If lesson is public, allow access for anyone
        elif is_public:
            pass  # Allow access - lesson is public
        # Otherwise deny access
        else:
            logger.info(f"Access denied for user {user_id} (role: {user_role}) to lesson {lesson_id} (teacher: {lesson_teacher_id}, public: {is_public})")
            return jsonify({'error': 'Access denied'}), 403
        
        return jsonify({
            'success': True,
            'lesson': lesson
        })
    except Exception as e:
        logger.error(f"Error getting lesson: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to get lesson: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>/view', methods=['GET'])
@login_required
def view_lesson(lesson_id):
    """Get a lesson with its versions for viewing"""
    try:
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        # Check if user can access this lesson
        user_role = session.get('role', 'student')
        user_id = session.get('user_id')
        lesson_teacher_id = lesson.get('teacher_id')
        is_public = lesson.get('is_public', False)
        
        # Users can access their own lessons (regardless of role)
        if lesson_teacher_id == user_id:
            pass  # Allow access - user owns this lesson
        # If lesson is public, allow access for anyone
        elif is_public:
            pass  # Allow access - lesson is public
        # Otherwise deny access
        else:
            logger.info(f"Access denied for user {user_id} (role: {user_role}) to lesson {lesson_id} (teacher: {lesson_teacher_id}, public: {is_public})")
            return jsonify({'error': 'Access denied'}), 403
        
        # Get lesson versions
        versions = LessonModel.get_lesson_versions(lesson_id)
        
        return jsonify({
            'success': True,
            'lesson': lesson,
            'versions': versions
        })
    except Exception as e:
        logger.error(f"Error viewing lesson: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to view lesson: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>', methods=['PUT'])
@teacher_required
def update_lesson(lesson_id):
    """Update a lesson (teacher only, and only their own lessons)"""
    try:
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        # Check if the lesson belongs to the current teacher
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        data = request.get_json()
        
        # Check if new title already exists for this teacher (if title is being changed)
        new_title = data.get('title')
        if new_title and new_title != lesson['title'] and LessonModel.check_title_exists(session['user_id'], new_title, exclude_lesson_id=lesson_id):
            return jsonify({'error': 'This lesson title is already used. Please choose a different title.'}), 400
        
        lesson_model = LessonModel(lesson_id)
        
        # Update lesson - if content is provided, also update original_content if it's empty
        content_to_save = data.get('content')
        if content_to_save:
            # Check if original_content is empty, if so, set it to the content being saved
            lesson = LessonModel.get_lesson_by_id(lesson_id)
            if lesson and (not lesson.get('original_content') or lesson.get('original_content').strip() == ''):
                # Update original_content as well if it's empty using ORM
                db = get_db()
                lesson_obj = db.query(DBLesson).filter(DBLesson.id == lesson_id).first()
                if lesson_obj:
                    lesson_obj.original_content = content_to_save
                    db.commit()
                    logger.info(f"Updated original_content for lesson {lesson_id} since it was empty")
        
        success = lesson_model.update_lesson(
            title=data.get('title'),
            summary=data.get('summary'),
            learning_objectives=data.get('learning_objectives'),
            focus_area=data.get('focus_area'),
            grade_level=data.get('grade_level'),
            content=content_to_save,
            is_public=data.get('is_public')
        )
        
        if success:
            return jsonify({
                'success': True,
                'message': 'Lesson updated successfully'
            })
        else:
            return jsonify({'error': 'Failed to update lesson'}), 500
            
    except Exception as e:
        logger.error(f"Error updating lesson: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to update lesson: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>', methods=['DELETE'])
@teacher_required
def delete_lesson(lesson_id):
    """Delete a lesson (teacher only, and only their own lessons)"""
    try:
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        # Check if the lesson belongs to the current teacher
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        lesson_model = LessonModel(lesson_id)
        success = lesson_model.delete_lesson()
        
        if success:
            return jsonify({
                'success': True,
                'message': 'Lesson deleted successfully'
            })
        else:
            return jsonify({'error': 'Failed to delete lesson'}), 500
            
    except Exception as e:
        logger.error(f"Error deleting lesson: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to delete lesson: {str(e)}'}), 500

@bp.route('/download_lesson/<int:lesson_id>', methods=['GET'])
@login_required
def download_lesson(lesson_id):
    """Download a lesson as DOCX file"""
    try:
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        # Check if user can access this lesson
        user_role = session.get('role', 'student')
        user_id = session.get('user_id')
        lesson_teacher_id = lesson.get('teacher_id')
        is_public = lesson.get('is_public', False)
        
        # Users can access their own lessons (regardless of role)
        if lesson_teacher_id == user_id:
            pass  # Allow access - user owns this lesson
        # If lesson is public, allow access for anyone
        elif is_public:
            pass  # Allow access - lesson is public
        # Otherwise deny access
        else:
            logger.info(f"Access denied for user {user_id} (role: {user_role}) to lesson {lesson_id} (teacher: {lesson_teacher_id}, public: {is_public})")
            return jsonify({'error': 'Access denied'}), 403
        
        # Create lesson structure for DOCX generation
        lesson_data = {
            'title': lesson['title'],
            'summary': lesson['summary'] or '',
            'learning_objectives': [lesson['learning_objectives']] if lesson['learning_objectives'] else [],
            'sections': [{'heading': 'Lesson Content', 'content': lesson['content']}],
            'key_concepts': [],
            'activities': [],
            'quiz': []
        }
        
        # Get API key from session
        api_key = session.get('groq_api_key')
        if not api_key:
            return jsonify({'error': 'API key not configured. Please set your API key first.'}), 400
        
        # Generate DOCX
        lesson_service = LessonService(api_key=api_key)
        docx_bytes = lesson_service._create_docx(lesson_data)
        
        # Create filename
        filename = lesson['title'].replace(' ', '_') + '.docx'
        
        # Create BytesIO object
        docx_buffer = BytesIO(docx_bytes)
        docx_buffer.seek(0)
        
        # Delete FAISS index after successful download
        try:
            lesson_service = LessonService(api_key=api_key)
            lesson_service._delete_faiss_index(lesson_id)
            logger.info(f"Deleted FAISS index after download for lesson {lesson_id}")
        except Exception as e:
            logger.warning(f"Failed to delete FAISS index for lesson {lesson_id}: {str(e)}")
        
        return send_file(
            docx_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Download error: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to download lesson: {str(e)}'}), 500 

@bp.route('/download_lesson_ppt/<int:lesson_id>', methods=['GET'])
@login_required
def download_lesson_ppt(lesson_id):
    """Download lesson as PowerPoint presentation"""
    lesson = LessonModel.get_lesson_by_id(lesson_id)
    if not lesson:
        return jsonify({'error': 'Lesson not found'}), 404

    # Check if user can access this lesson
    user_role = session.get('role', 'student')
    user_id = session.get('user_id')
    lesson_teacher_id = lesson.get('teacher_id')
    is_public = lesson.get('is_public', False)
    
    # Users can access their own lessons (regardless of role)
    if lesson_teacher_id == user_id:
        pass  # Allow access - user owns this lesson
    # If lesson is public, allow access for anyone
    elif is_public:
        pass  # Allow access - lesson is public
    # Otherwise deny access
    else:
        logger.info(f"Access denied for user {user_id} (role: {user_role}) to lesson {lesson_id} (teacher: {lesson_teacher_id}, public: {is_public})")
        return jsonify({'error': 'Access denied'}), 403

    # Prepare lesson data for PPT generation
    lesson_data = {
        'title': lesson['title'],
        'summary': lesson['summary'] or '',
        'learning_objectives': [lesson['learning_objectives']] if lesson['learning_objectives'] else [],
        'sections': [{'heading': 'Lesson Content', 'content': lesson['content']}],
        'key_concepts': [],
        'activities': [],
        'quiz': []
    }
    
    # Get API key from session
    api_key = session.get('groq_api_key')
    if not api_key:
        return jsonify({'error': 'API key not configured. Please set your API key first.'}), 400
    
    try:
        lesson_service = LessonService(api_key=api_key)
        ppt_bytes = lesson_service.create_ppt(lesson_data)
        
        # Create filename
        filename = lesson['title'].replace(' ', '_') + '.pptx'
        
        # Create response
        from io import BytesIO
        ppt_buffer = BytesIO(ppt_bytes)
        ppt_buffer.seek(0)
        
        # Delete FAISS index after successful download
        try:
            lesson_service = LessonService(api_key=api_key)
            lesson_service._delete_faiss_index(lesson_id)
            logger.info(f"Deleted FAISS index after PPT download for lesson {lesson_id}")
        except Exception as e:
            logger.warning(f"Failed to delete FAISS index for lesson {lesson_id}: {str(e)}")
        
        return send_file(
            ppt_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.presentationml.presentation'
        )
    except Exception as e:
        logger.error(f"PPT generation error: {str(e)}")
        return jsonify({'error': f'Failed to generate PPT: {str(e)}'}), 500

@bp.route('/download_lesson_pdf/<int:lesson_id>', methods=['GET'])
@login_required
def download_lesson_pdf(lesson_id):
    """Download a lesson as PDF file"""
    try:
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404

        # Check if user can access this lesson
        user_role = session.get('role', 'student')
        user_id = session.get('user_id')
        lesson_teacher_id = lesson.get('teacher_id')
        is_public = lesson.get('is_public', False)
        
        # Users can access their own lessons (regardless of role)
        if lesson_teacher_id == user_id:
            pass  # Allow access - user owns this lesson
        # If lesson is public, allow access for anyone
        elif is_public:
            pass  # Allow access - lesson is public
        # Otherwise deny access
        else:
            logger.info(f"Access denied for user {user_id} (role: {user_role}) to lesson {lesson_id} (teacher: {lesson_teacher_id}, public: {is_public})")
            return jsonify({'error': 'Access denied'}), 403

        # Try to generate PDF using reportlab (preferred method)
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
            
            # Create PDF buffer
            pdf_buffer = BytesIO()
            
            # Create PDF document
            doc = SimpleDocTemplate(pdf_buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            
            # Define custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=12,
                spaceBefore=20
            )
            
            content_style = ParagraphStyle(
                'CustomContent',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                alignment=TA_LEFT
            )
            
            # Build PDF content
            story = []
            
            # Add title
            story.append(Paragraph(lesson['title'], title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Add summary if available
            if lesson.get('summary'):
                story.append(Paragraph("Summary", heading_style))
                story.append(Paragraph(lesson['summary'], content_style))
                story.append(Spacer(1, 0.1*inch))
            
            # Add learning objectives if available
            if lesson.get('learning_objectives'):
                story.append(Paragraph("Learning Objectives", heading_style))
                story.append(Paragraph(lesson['learning_objectives'], content_style))
                story.append(Spacer(1, 0.1*inch))
            
            # Add focus area and grade level
            if lesson.get('focus_area'):
                story.append(Paragraph("Focus Area", heading_style))
                story.append(Paragraph(lesson['focus_area'], content_style))
                story.append(Spacer(1, 0.1*inch))
            
            if lesson.get('grade_level'):
                story.append(Paragraph("Grade Level", heading_style))
                story.append(Paragraph(lesson['grade_level'], content_style))
                story.append(Spacer(1, 0.1*inch))
            
            # Add main content
            if lesson.get('content'):
                story.append(Paragraph("Lesson Content", heading_style))
                # Split content by paragraphs and create Paragraph objects
                content_paragraphs = lesson['content'].split('\n\n')
                for para in content_paragraphs:
                    if para.strip():
                        story.append(Paragraph(para.strip(), content_style))
                        story.append(Spacer(1, 0.05*inch))
            
            # Build PDF
            doc.build(story)
            pdf_buffer.seek(0)
            
            # Create filename
            filename = lesson['title'].replace(' ', '_').replace('/', '_') + '.pdf'
            
            # Delete FAISS index after successful download
            try:
                lesson_service = LessonService(api_key=api_key)
                lesson_service._delete_faiss_index(lesson_id)
                logger.info(f"Deleted FAISS index after PDF download for lesson {lesson_id}")
            except Exception as e:
                logger.warning(f"Failed to delete FAISS index for lesson {lesson_id}: {str(e)}")
            
            return send_file(
                pdf_buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='application/pdf'
            )
            
        except ImportError:
            logger.warning("ReportLab not available, falling back to LibreOffice method")
            # Fallback to LibreOffice method if reportlab is not available
            pass

        # Fallback: Generate DOCX first and convert with LibreOffice
        lesson_data = {
            'title': lesson['title'],
            'summary': lesson['summary'] or '',
            'learning_objectives': [lesson['learning_objectives']] if lesson['learning_objectives'] else [],
            'sections': [{'heading': 'Lesson Content', 'content': lesson['content']}],
            'key_concepts': [],
            'activities': [],
            'quiz': []
        }
        
        # Get API key from session
        api_key = session.get('groq_api_key')
        if not api_key:
            return jsonify({'error': 'API key not configured. Please set your API key first.'}), 400
        
        lesson_service = LessonService(api_key=api_key)
        docx_bytes = lesson_service._create_docx(lesson_data)

        # Save DOCX to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as docx_file:
            docx_file.write(docx_bytes)
            docx_path = docx_file.name

        pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
        try:
            import subprocess
            # Use LibreOffice to convert DOCX to PDF
            result = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), docx_path
            ], capture_output=True, timeout=30)
            
            if result.returncode != 0:
                raise Exception(f'LibreOffice conversion failed: {result.stderr.decode()}')

            @after_this_request
            def cleanup(response):
                try:
                    os.remove(docx_path)
                except Exception:
                    pass
                try:
                    os.remove(pdf_path)
                except Exception:
                    pass
                return response

            # Delete FAISS index after successful download
            try:
                lesson_service = LessonService(api_key=api_key)
                lesson_service._delete_faiss_index(lesson_id)
                logger.info(f"Deleted FAISS index after PDF download for lesson {lesson_id}")
            except Exception as e:
                logger.warning(f"Failed to delete FAISS index for lesson {lesson_id}: {str(e)}")
            
            return send_file(pdf_path, as_attachment=True, download_name=lesson['title'].replace('/', '_') + '.pdf', mimetype='application/pdf')
            
        except Exception as e:
            # Clean up files on error
            if os.path.exists(docx_path):
                os.remove(docx_path)
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
            logger.error(f"LibreOffice PDF generation error: {str(e)}")
            return jsonify({'error': 'PDF generation failed. Please try downloading as DOCX instead.'}), 500
            
    except Exception as e:
        logger.error(f"PDF generation error: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to generate PDF: {str(e)}'}), 500

@bp.route('/ask_question', methods=['POST'])
@student_required
def ask_lesson_question():
    data = request.get_json()
    lesson_id = data.get('lesson_id')
    question = data.get('question')
    if not lesson_id or not question:
        return jsonify({'error': 'lesson_id and question are required'}), 400
    
    api_key = session.get('groq_api_key')
    if not api_key:
        return jsonify({'error': 'API key not configured. Please set your API key first.'}), 400
    
    service = LessonService(api_key=api_key)
    
    # Get conversation history for context
    from app.models.models import LessonChatHistory, LessonFAQ
    user_id = session['user_id']
    conversation_history = LessonChatHistory.get_lesson_chat_history(lesson_id, user_id)
    
    # Use the new query analysis system for better responses
    query_analysis = service.analyze_user_query(question)
    logger.info(f"Lesson question analysis: {query_analysis}")
    
    # Try to answer using the specific lesson first with conversation context
    result = service.answer_lesson_question(lesson_id, question, conversation_history)
    
    # If the specific lesson doesn't have good results, try general search with context
    if 'error' in result or not result.get('answer') or len(result.get('answer', '')) < 50:
        logger.info("Specific lesson search didn't provide good results, trying general search")
        
        # Get available lesson IDs for broader context
        available_lessons = service._get_available_lesson_ids()
        if available_lessons:
            general_result = service.answer_general_question(question, available_lessons, conversation_history)
            if general_result.get('answer') and len(general_result.get('answer', '')) > 50:
                result = general_result
                logger.info("Using general search result for better answer")
    
    if 'error' in result:
        return jsonify({'error': result['error']}), 400
    
    # Save the Q&A to lesson chat history
    canonical = result.get('canonical_question', question)
    LessonChatHistory.save_qa(lesson_id, user_id, question, result['answer'], canonical_question=canonical)
    
    # Log the question to FAQ table for teacher visibility
    try:
        LessonFAQ.log_question(lesson_id, canonical)
        logger.info(f"Question logged to FAQ table for lesson {lesson_id}: {canonical}")
    except Exception as e:
        logger.error(f"Error logging question to FAQ table: {str(e)}")
    
    return jsonify({
        'answer': result['answer'], 
        'canonical_question': canonical,
        'source': result.get('source', 'unknown'),
        'confidence': result.get('confidence', 0.8),
        'query_analysis': query_analysis
    })

@bp.route('/lesson_chat_history/<int:lesson_id>', methods=['GET'])
@login_required
def get_lesson_chat_history(lesson_id):
    """Get chat history for a specific lesson"""
    try:
        from app.models.models import LessonChatHistory
        user_id = session['user_id']
        history = LessonChatHistory.get_lesson_chat_history(lesson_id, user_id)
        return jsonify({'history': history})
    except Exception as e:
        logger.error(f"Error getting lesson chat history: {str(e)}")
        return jsonify({'error': 'Failed to get lesson chat history'}), 500

@bp.route('/clear_lesson_chat_history/<int:lesson_id>', methods=['DELETE'])
@login_required
def clear_lesson_chat_history(lesson_id):
    """Clear chat history for a specific lesson"""
    try:
        from app.models.models import LessonChatHistory
        user_id = session['user_id']
        LessonChatHistory.clear_lesson_chat_history(lesson_id, user_id)
        return jsonify({'message': 'Lesson chat history cleared successfully'})
    except Exception as e:
        logger.error(f"Error clearing lesson chat history: {str(e)}")
        return jsonify({'error': 'Failed to clear lesson chat history'}), 500

@bp.route('/lesson/<int:lesson_id>/create_version', methods=['POST'])
@teacher_required
def create_lesson_version(lesson_id):
    """Create a new version of an existing lesson"""
    try:
        # Check if original lesson exists and belongs to the teacher
        original_lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not original_lesson:
            return jsonify({'error': 'Original lesson not found'}), 404
        
        if original_lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        # Keep source_version_id as the lesson we're branching from (could be original or a child)
        source_version_id = lesson_id
        
        # For display defaults, load the root/original lesson's data if needed
        root_lesson_id = original_lesson.get('parent_lesson_id') or lesson_id
        original_lesson_data = LessonModel.get_lesson_by_id(root_lesson_id) or original_lesson
        
        data = request.get_json()
        
        # Debug logging
        logger.info(f"DEBUG: create_lesson_version - received data: {data}")
        logger.info(f"DEBUG: create_lesson_version - content field: {data.get('content', 'NOT_PROVIDED')}")
        logger.info(f"DEBUG: create_lesson_version - content length: {len(data.get('content', ''))}")
        
        # Check if new title already exists for this teacher (if title is being changed)
        new_title = data.get('title', original_lesson_data['title'])
        if new_title != original_lesson_data['title'] and LessonModel.check_title_exists(session['user_id'], new_title):
            return jsonify({'error': 'This lesson title is already used. Please choose a different title.'}), 400
        
        # Create new version: pass source_version_id; model will attach to root and flag the source
        new_lesson_id = LessonModel.create_new_version(
            original_lesson_id=source_version_id,
            teacher_id=session['user_id'],
            title=data.get('title', original_lesson_data['title']),
            summary=data.get('summary', original_lesson_data['summary']),
            learning_objectives=data.get('learning_objectives', original_lesson_data['learning_objectives']),
            focus_area=data.get('focus_area', original_lesson_data['focus_area']),
            grade_level=data.get('grade_level', original_lesson_data['grade_level']),
            content=data.get('content', original_lesson_data['content']),
            file_name=data.get('file_name', original_lesson_data.get('file_name'))
        )
        
        return jsonify({
            'success': True,
            'new_lesson_id': new_lesson_id,
            'message': 'New version created successfully'
        })
        
    except Exception as e:
        logger.error(f"Error creating lesson version: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to create lesson version: {str(e)}'}), 500

@bp.route('/check_title_exists', methods=['GET'])
@teacher_required
def check_title_exists():
    """Check if a lesson title already exists for the current teacher"""
    try:
        title = request.args.get('title', '').strip()
        if not title:
            return jsonify({'exists': False})
        
        teacher_id = session['user_id']
        exists = LessonModel.check_title_exists(teacher_id, title)
        
        return jsonify({'exists': exists})
        
    except Exception as e:
        logger.error(f"Error checking title existence: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to check title: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>/create_ai_version', methods=['POST'])
@teacher_required
def create_ai_lesson_version(lesson_id):
    """Create an AI-improved version of an existing lesson"""
    try:
        # Check if original lesson exists and belongs to the teacher
        original_lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not original_lesson:
            return jsonify({'error': 'Original lesson not found'}), 404
        
        if original_lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        # Keep source_version_id as the lesson we're branching from (could be original or a child)
        source_version_id = lesson_id
        
        # For display defaults, load the root/original lesson's data if needed
        root_lesson_id = original_lesson.get('parent_lesson_id') or lesson_id
        original_lesson_data = LessonModel.get_lesson_by_id(root_lesson_id) or original_lesson
        
        data = request.get_json()
        improvement_prompt = data.get('improvement_prompt', '')
        
        # Get API key from session
        api_key = session.get('groq_api_key')
        if not api_key:
            return jsonify({'error': 'API key not configured. Please set your API key first.'}), 400
        
        # Use LessonService to improve the lesson
        lesson_service = LessonService(api_key=api_key)
        
        # Generate improved lesson content
        improved_content = lesson_service.improve_lesson_content(
            lesson_id=lesson_id,
            current_content=original_lesson_data['content'],
            improvement_prompt=improvement_prompt
        )
        
        # Create new version with improved content, always using the original lesson ID
        new_lesson_id = LessonModel.create_new_version(
            original_lesson_id=lesson_id,
            teacher_id=session['user_id'],
            title=original_lesson_data['title'],
            summary=original_lesson_data['summary'],
            learning_objectives=original_lesson_data['learning_objectives'],
            focus_area=original_lesson_data['focus_area'],
            grade_level=original_lesson_data['grade_level'],
            content=improved_content,
            file_name=original_lesson_data.get('file_name')
        )
        
        return jsonify({
            'success': True,
            'new_lesson_id': new_lesson_id,
            'message': 'AI-improved version created successfully',
            'improved_lesson': {
                'title': original_lesson_data['title'],
                'summary': original_lesson_data['summary'],
                'learning_objectives': original_lesson_data['learning_objectives'],
                'focus_area': original_lesson_data['focus_area'],
                'grade_level': original_lesson_data['grade_level'],
                'content': improved_content
            }
        })
        
    except Exception as e:
        logger.error(f"Error creating AI lesson version: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to create AI lesson version: {str(e)}'}), 500

@bp.route('/faqs/<int:lesson_id>', methods=['GET'])
@teacher_required
def get_lesson_faqs(lesson_id):
    try:
        # Get lesson details first
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        # Read FAQs from the lesson_faq table (real student questions)
        import sqlite3
        conn = sqlite3.connect('instance/chatbot.db')
        c = conn.cursor()
        
        # Get questions from lesson_faq table for this specific lesson
        # Use canonical form when available
        c.execute('SELECT COALESCE(canonical_question, question) as question, count FROM lesson_faq WHERE lesson_id=? ORDER BY count DESC', (lesson_id,))
        faq_rows = c.fetchall()
        conn.close()
        
        # Format the FAQs to match the expected structure
        faqs = []
        
        # Prepare DB for fetching latest answers from lesson_chat_history
        import sqlite3 as _sqlite3
        _conn2 = _sqlite3.connect('instance/chatbot.db')
        _c2 = _conn2.cursor()
        # Ensure table exists (defensive)
        _c2.execute('''CREATE TABLE IF NOT EXISTS lesson_chat_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            lesson_id INTEGER,
            user_id INTEGER,
            question TEXT,
            answer TEXT,
            canonical_question TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )''')

        for row in faq_rows:
            # Determine version display text
            version_text = ""
            if lesson.get('parent_lesson_id'):
                version_text = f"v{lesson.get('version', 1)}"
            else:
                version_text = "v1 (Original)"
            
            canonical_or_question = row[0]
            # Fetch latest answer using canonical match first, then exact, then partial
            latest_answer = None
            try:
                # Try canonical match first
                _c2.execute(
                    'SELECT answer FROM lesson_chat_history WHERE lesson_id = ? AND canonical_question = ? ORDER BY datetime(created_at) DESC LIMIT 1',
                    (lesson_id, canonical_or_question)
                )
                _row_ans = _c2.fetchone()
                if not _row_ans:
                    # Fallback exact text match
                    _c2.execute(
                        'SELECT answer FROM lesson_chat_history WHERE lesson_id = ? AND question = ? ORDER BY datetime(created_at) DESC LIMIT 1',
                        (lesson_id, canonical_or_question)
                    )
                    _row_ans = _c2.fetchone()
                if not _row_ans:
                    # Fallback: partial match using LIKE
                    _c2.execute(
                        'SELECT answer FROM lesson_chat_history WHERE lesson_id = ? AND (question LIKE ? OR canonical_question LIKE ?) ORDER BY datetime(created_at) DESC LIMIT 1',
                        (lesson_id, f"%{canonical_or_question[:30]}%", f"%{canonical_or_question[:30]}%")
                    )
                    _row_ans = _c2.fetchone()
                if _row_ans:
                    latest_answer = _row_ans[0]
            except Exception:
                latest_answer = None

            faqs.append({
                'question': row[0],
                'count': row[1],
                'times_asked': row[1],  # For compatibility with frontend
                'lessonTitle': f"{lesson.get('title', '')} {version_text}",
                'subject': lesson.get('focus_area', ''),
                'grade': lesson.get('grade_level', ''),
                'time_ago': 'Recently',  # Placeholder since timestamps aren't tracked
                'version': lesson.get('version', 1),
                'is_version': lesson.get('parent_lesson_id') is not None,
                'parent_lesson_id': lesson.get('parent_lesson_id'),
                'answer': latest_answer
            })
        _conn2.close()
        
        return jsonify({'faqs': faqs})
        
    except Exception as e:
        logger.error(f"Error getting lesson FAQs: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to load lesson FAQs'}), 500

@bp.route('/faqs_count/<int:lesson_id>', methods=['GET'])
@login_required
def get_lesson_faq_count(lesson_id):
    try:
        import sqlite3
        conn = sqlite3.connect('instance/chatbot.db')
        c = conn.cursor()
        c.execute('''CREATE TABLE IF NOT EXISTS lesson_faq (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            lesson_id INTEGER,
            question TEXT,
            count INTEGER DEFAULT 1,
            canonical_question TEXT
        )''')
        c.execute('SELECT COALESCE(SUM(count), 0) FROM lesson_faq WHERE lesson_id=?', (lesson_id,))
        total = c.fetchone()[0] or 0
        c.execute('SELECT COUNT(*) FROM lesson_faq WHERE lesson_id=?', (lesson_id,))
        unique_qs = c.fetchone()[0] or 0
        conn.close()
        return jsonify({'total_count': int(total), 'unique_count': int(unique_qs)})
    except Exception:
        return jsonify({'total_count': 0, 'unique_count': 0})

@bp.route('/faq_dashboard', methods=['GET'])
@teacher_required
def faq_dashboard():
    try:
        user_id = session['user_id']
        # Get all lessons for this teacher
        lessons = LessonModel.get_lessons_by_teacher(user_id)
        lesson_ids = [lesson['id'] for lesson in lessons]
        top_questions = []
        total_questions = 0
        recent_questions = []
        import sqlite3
        from datetime import datetime, timedelta

        conn = sqlite3.connect('instance/chatbot.db')
        c = conn.cursor()
        # For each lesson, get top questions
        for lesson in lessons:
            c.execute('SELECT COALESCE(canonical_question, question) as question, count FROM lesson_faq WHERE lesson_id=? ORDER BY count DESC LIMIT 3', (lesson['id'],))
            faqs = [{'question': row[0], 'count': row[1]} for row in c.fetchall()]
            total_questions += sum(row['count'] for row in faqs)
            if faqs:
                top_questions.append({
                    'title': lesson['title'],
                    'subject': lesson.get('focus_area', ''),
                    'questions': faqs
                })
        # Get recent questions (last 24h) -- placeholder, as timestamps are not tracked
        # If you want real recent questions, you need to log timestamps in lesson_faq
        # For now, just return the most asked question per lesson
        for lesson in lessons:
            c.execute('SELECT COALESCE(canonical_question, question) as question, count FROM lesson_faq WHERE lesson_id=? ORDER BY count DESC LIMIT 1', (lesson['id'],))
            row = c.fetchone()
            if row:
                recent_questions.append({
                    'question': row[0],
                    'student_name': '',  # Not tracked
                    'lesson_title': lesson['title'],
                    'time_ago': 'Recently'
                })
        conn.close()
        return jsonify({
            'total_questions': total_questions,
            'weekly_questions': total_questions,  # Placeholder
            'active_students': 0,  # Placeholder
            'top_questions': top_questions,
            'recent_questions': recent_questions
        })
    except Exception as e:
        logger.error(f"Error loading FAQ dashboard: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to load FAQ dashboard'}), 500 

@bp.route('/export_faq_dashboard', methods=['GET'])
@teacher_required
def export_faq_dashboard():
    try:
        user_id = session['user_id']
        lessons = LessonModel.get_lessons_by_teacher(user_id)
        import sqlite3
        from io import BytesIO
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        conn = sqlite3.connect('instance/chatbot.db')
        c = conn.cursor()
        # Prepare data for export
        rows = []
        for lesson in lessons:
            c.execute('SELECT COALESCE(canonical_question, question) as question, count FROM lesson_faq WHERE lesson_id=? ORDER BY count DESC', (lesson['id'],))
            faqs = c.fetchall()
            for faq in faqs:
                rows.append({
                    'lesson_title': lesson['title'],
                    'subject': lesson.get('focus_area', ''),
                    'question': faq[0],
                    'count': faq[1]
                })
        conn.close()

        # Create Word document
        doc = Document()
        doc.add_heading('FAQ Dashboard Export', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Lesson Title'
        hdr_cells[1].text = 'Subject'
        hdr_cells[2].text = 'Question'
        hdr_cells[3].text = 'Count'
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.runs[0].font.bold = True
                paragraph.runs[0].font.size = Pt(11)
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = str(row['lesson_title'])
            cells[1].text = str(row['subject'])
            cells[2].text = str(row['question'])
            cells[3].text = str(row['count'])
        # Save to BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        from flask import make_response
        response = make_response(doc_io.getvalue())
        response.headers["Content-Disposition"] = "attachment; filename=faq_dashboard_export.docx"
        response.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return response
    except Exception as e:
        logger.error(f"Error exporting FAQ dashboard: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to export FAQ dashboard'}), 500

@bp.route('/lesson/<int:lesson_id>/save_draft', methods=['POST'])
@teacher_required
def save_lesson_draft(lesson_id):
    """Save draft content for a lesson"""
    try:
        # Check if lesson exists and belongs to the teacher
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        data = request.get_json()
        draft_content = data.get('draft_content', '')
        
        success = LessonModel.save_draft_content(lesson_id, draft_content)
        
        if success:
            return jsonify({
                'success': True,
                'message': 'Draft content saved successfully'
            })
        else:
            return jsonify({'error': 'Failed to save draft content'}), 500
            
    except Exception as e:
        logger.error(f"Error saving draft content: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to save draft content: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>/get_draft', methods=['GET'])
@teacher_required
def get_lesson_draft(lesson_id):
    """Get draft content for a lesson"""
    try:
        # Check if lesson exists and belongs to the teacher
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        draft_content = LessonModel.get_draft_content(lesson_id)
        
        # Get the current content from the lesson (this is what we update when finalizing)
        current_content = lesson.get('content') or lesson.get('original_content') or ''
        
        # Get the original content from the first version of this lesson
        original_content = lesson.get('original_content', '')
        if lesson.get('lesson_id'):
            # Get the first version (version_number = 1) to get true original content using ORM
            db = get_db()
            first_version = db.query(DBLesson).filter(
                DBLesson.lesson_id == lesson['lesson_id'],
                DBLesson.version_number == 1
            ).first()
            if first_version:
                # Use content if original_content is empty, otherwise use original_content
                original_content = first_version.original_content or first_version.content or ''
        
        # Log for debugging
        logger.info(f"get_draft for lesson {lesson_id}: current_content length={len(current_content)}, original_content length={len(original_content)}")
        
        return jsonify({
            'success': True,
            'draft_content': draft_content,
            'current_content': current_content,  # Current version's content
            'original_content': original_content   # True original content from first version
        })
            
    except Exception as e:
        logger.error(f"Error getting draft content: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to get draft content: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>/review_by_ai', methods=['POST'])
@teacher_required
def review_by_ai(lesson_id):
    """Review lesson content by AI based on user query - provides fresh response without showing previous content"""
    try:
        # Check if lesson exists and belongs to the teacher
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        data = request.get_json()
        user_query = data.get('query', '')
        
        if not user_query.strip():
            return jsonify({'error': 'Query is required'}), 400
        
        # Get API key from session
        api_key = session.get('groq_api_key')
        if not api_key:
            return jsonify({'error': 'API key not configured. Please set your API key first.'}), 400
        
        # Use LessonService to provide AI review
        lesson_service = LessonService(api_key=api_key)
        ai_response = lesson_service.review_by_ai(
            lesson_id=lesson_id,
            user_query=user_query
        )
        
        return jsonify({
            'success': True,
            'ai_response': ai_response,
            'message': 'AI review completed successfully'
        })
            
    except Exception as e:
        logger.error(f"Error getting AI review: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to get AI review: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>/apply_prompt', methods=['POST'])
@teacher_required
def apply_prompt_to_lesson(lesson_id):
    """Apply AI prompt to lesson content and save as draft"""
    try:
        # Check if lesson exists and belongs to the teacher
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        data = request.get_json()
        prompt = data.get('prompt', '')
        
        if not prompt.strip():
            return jsonify({'error': 'Prompt is required'}), 400
        
        # Get API key from session
        api_key = session.get('groq_api_key')
        if not api_key:
            return jsonify({'error': 'API key not configured. Please set your API key first.'}), 400
        
        # Get current content (use draft if available, otherwise original)
        current_draft = LessonModel.get_draft_content(lesson_id)
        content_to_edit = current_draft if current_draft else lesson.get('original_content', lesson['content'])
        
        # Use LessonService to apply the prompt
        lesson_service = LessonService(api_key=api_key)
        improved_content = lesson_service.improve_lesson_content(
            lesson_id=lesson_id,
            current_content=content_to_edit,
            improvement_prompt=prompt
        )
        
        # Check if the improved content is a JSON string and extract the actual content
        import json
        
        # Recursive function to extract text content from JSON
        def extract_content(data):
            if isinstance(data, str):
                # Try to parse as JSON
                try:
                    parsed = json.loads(data)
                    return extract_content(parsed)
                except (json.JSONDecodeError, TypeError):
                    return data
            elif isinstance(data, dict):
                # Look for content in various fields
                for key in ['answer', 'response', 'content', 'text', 'message']:
                    if key in data:
                        return extract_content(data[key])
                # If no content field found, return string representation
                return str(data)
            elif isinstance(data, list):
                # Join list items
                return '\n'.join([extract_content(item) for item in data])
            else:
                return str(data)
        
        improved_content = extract_content(improved_content)
        
        # Save the improved content as draft
        success = LessonModel.save_draft_content(lesson_id, improved_content)
        
        if success:
            return jsonify({
                'success': True,
                'draft_content': improved_content,
                'message': 'Prompt applied successfully'
            })
        else:
            return jsonify({'error': 'Failed to save draft content'}), 500
            
    except Exception as e:
        logger.error(f"Error applying prompt: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to apply prompt: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>/finalize_version', methods=['POST'])
@teacher_required
def finalize_lesson_version(lesson_id):
    """Finalize draft content into a new lesson version"""
    try:
        # Check if lesson exists and belongs to the teacher
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        # Get draft content
        draft_content = LessonModel.get_draft_content(lesson_id)
        
        if not draft_content.strip():
            return jsonify({'error': 'No draft content to finalize. Apply a prompt first.'}), 400
        
        # Extract actual text content from JSON if needed
        import json
        
        def extract_content(data):
            if isinstance(data, str):
                try:
                    parsed = json.loads(data)
                    return extract_content(parsed)
                except (json.JSONDecodeError, TypeError):
                    return data
            elif isinstance(data, dict):
                for key in ['answer', 'response', 'content', 'text', 'message']:
                    if key in data:
                        return extract_content(data[key])
                return str(data)
            elif isinstance(data, list):
                return '\n'.join([extract_content(item) for item in data])
            else:
                return str(data)
        
        # Extract the actual content from JSON if the draft is in JSON format
        draft_content = extract_content(draft_content)
        
        # Determine the actual original lesson ID
        original_lesson_id = lesson.get('parent_lesson_id') or lesson_id
        
        # Create new version using draft content: pass source_version_id; model will attach to root and flag the source
        new_lesson_id = LessonModel.create_new_version_from_draft(
            original_lesson_id=lesson_id,
            teacher_id=session['user_id'],
            title=lesson['title'],
            summary=lesson['summary'],
            learning_objectives=lesson['learning_objectives'],
            focus_area=lesson['focus_area'],
            grade_level=lesson['grade_level'],
            draft_content=draft_content,
            file_name=lesson.get('file_name'),
            is_public=lesson.get('is_public', True)
        )
        
        # Update the new lesson with the final content in database
        new_lesson_model = LessonModel(new_lesson_id)
        new_lesson_model.update_lesson(content=draft_content)
        
        # Delete FAISS index after storing in database
        try:
            lesson_service = LessonService(api_key=session.get('groq_api_key'))
            lesson_service._delete_faiss_index(new_lesson_id)
            logger.info(f"Deleted FAISS index for finalized lesson {new_lesson_id}")
        except Exception as e:
            logger.warning(f"Failed to delete FAISS index for lesson {new_lesson_id}: {str(e)}")
        
        # Clear the draft content from the current lesson
        LessonModel.clear_draft_content(lesson_id)
        
        return jsonify({
            'success': True,
            'new_lesson_id': new_lesson_id,
            'message': 'New version created successfully'
        })
        
    except Exception as e:
        logger.error(f"Error finalizing version: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to finalize version: {str(e)}'}), 500


#interactive chat with lesson

# ROUTE HANDLER
@bp.route('/lesson/<int:lesson_id>/interactive_chat', methods=['POST'])
@teacher_required
def interactive_chat(lesson_id):
    try:
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403

        data = request.get_json()
        user_query = data.get('query', '')

        if not user_query.strip():
            return jsonify({'error': 'Query is required'}), 400

        api_key = session.get('groq_api_key')
        if not api_key:
            return jsonify({'error': 'API key not configured'}), 400

        # Get lesson data to pass form context
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404

        lesson_service = LessonService(api_key=api_key)
        
        # Check if document was uploaded (by checking if vector store exists)
        # For now, we'll assume document was uploaded if focus_area exists
        document_uploaded = bool(lesson.get('focus_area'))
        
        result = lesson_service.interactive_chat(
            lesson_id=lesson_id,
            user_query=user_query,
            subject=lesson.get('focus_area'),
            grade_level=lesson.get('grade_level'),
            focus_area=lesson.get('focus_area'),
            document_uploaded=document_uploaded,
            document_filename=None  # Can be enhanced to store filename
        )

        # If complete lesson is generated, save it to database
        if result.complete_lesson == "yes":
            # Update lesson content in database
            lesson_model = LessonModel(lesson_id)
            lesson_model.update_lesson(content=result.ai_response)
            logger.info(f"Complete lesson saved to database for lesson_id: {lesson_id}")

        # Return response with lesson update status
        return jsonify({
            'success': True,
            'ai_response': result.ai_response,
            'complete_lesson': result.complete_lesson,
            'lesson_id': lesson_id
        })

    except Exception as e:
        logger.error(f"Error in interactive chat: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to process chat: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>/interactive_chat_stream', methods=['POST'])
@teacher_required
def interactive_chat_stream(lesson_id):
    """Interactive chat with streaming response using Server-Sent Events (SSE)"""
    try:
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403

        data = request.get_json()
        user_query = data.get('query', '')

        if not user_query.strip():
            return jsonify({'error': 'Query is required'}), 400

        api_key = session.get('groq_api_key')
        if not api_key:
            return jsonify({'error': 'API key not configured'}), 400

        # Get lesson data to pass form context
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404

        lesson_service = LessonService(api_key=api_key)
        
        # Check if document was uploaded
        document_uploaded = bool(lesson.get('focus_area'))
        
        # Import Response for streaming
        from flask import Response
        import json
        
        def generate():
            """Generator function for Server-Sent Events"""
            try:
                full_response = ""
                complete_lesson_status = "no"
                
                # Stream the response
                for chunk_text, is_complete, complete_lesson in lesson_service.interactive_chat_stream(
                    lesson_id=lesson_id,
                    user_query=user_query,
                    subject=lesson.get('focus_area'),
                    grade_level=lesson.get('grade_level'),
                    focus_area=lesson.get('focus_area'),
                    document_uploaded=document_uploaded,
                    document_filename=None
                ):
                    # Always accumulate the chunk text (even if empty for final message)
                    if chunk_text:
                        full_response += chunk_text
                        # Send chunk as SSE
                        yield f"data: {json.dumps({'chunk': chunk_text, 'is_complete': False, 'complete_lesson': 'no'})}\n\n"
                    
                    # Check if this is the final message
                    if is_complete:
                        complete_lesson_status = complete_lesson
                        # Send completion message with full response
                        yield f"data: {json.dumps({'chunk': '', 'is_complete': True, 'complete_lesson': complete_lesson_status, 'full_response': full_response})}\n\n"
                        
                        # If complete lesson is generated, save it to database
                        if complete_lesson_status == "yes" and full_response:
                            try:
                                lesson_model = LessonModel(lesson_id)
                                lesson_model.update_lesson(content=full_response)
                                logger.info(f"Complete lesson saved to database for lesson_id: {lesson_id}")
                            except Exception as e:
                                logger.error(f"Error saving lesson: {str(e)}")
                        
                        break
                
            except Exception as e:
                logger.error(f"Error in streaming chat: {str(e)}", exc_info=True)
                error_data = json.dumps({
                    'error': str(e),
                    'is_complete': True,
                    'complete_lesson': 'no'
                })
                yield f"data: {error_data}\n\n"
        
        # Return streaming response with SSE headers
        return Response(
            generate(),
            mimetype='text/event-stream',
            headers={
                'Cache-Control': 'no-cache',
                'X-Accel-Buffering': 'no',  # Disable buffering in nginx
                'Connection': 'keep-alive'
            }
        )

    except Exception as e:
        logger.error(f"Error in streaming chat route: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to start streaming: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>/clear_draft', methods=['DELETE'])
@teacher_required
def clear_lesson_draft(lesson_id):
    """Clear draft content for a lesson"""
    try:
        # Check if lesson exists and belongs to the teacher
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        success = LessonModel.clear_draft_content(lesson_id)
        
        if success:
            return jsonify({
                'success': True,
                'message': 'Draft content cleared successfully'
            })
        else:
            return jsonify({'error': 'Failed to clear draft content'}), 500
            
    except Exception as e:
        logger.error(f"Error clearing draft content: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to clear draft content: {str(e)}'}), 500 


@bp.route('/chatbot', methods=['GET'])
@teacher_required
def chatbot():
    """Render the chatbot interface"""
    try:
        # Get user's lessons for selection
        user_id = session.get('user_id')
        lessons = LessonModel.get_lessons_by_teacher(user_id)
        
        return render_template('chatbot.html', lessons=lessons)
    except Exception as e:
        logger.error(f"Error rendering chatbot: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to render chatbot: {str(e)}'}), 500