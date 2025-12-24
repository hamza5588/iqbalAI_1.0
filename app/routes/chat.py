# app/routes/chat.py
from flask import Blueprint, redirect, request, session, jsonify, render_template, url_for, send_file
from app.services import ChatService, PromptService
from app.models.models import SurveyModel, LessonModel
# from app.utils.decorators import login_required
from app.utils.auth import login_required
from app.utils.decorators import teacher_required
import logging
from app.utils.db import get_db
import time
from functools import lru_cache
from io import BytesIO
from langdetect import detect
from gtts import gTTS
import os

from openai import OpenAI

logger = logging.getLogger(__name__)
bp = Blueprint('chat', __name__)

# Simple cache for token status (user_id -> (timestamp, data))
_token_status_cache = {}
CACHE_TTL = 2  # Cache for 2 seconds

@bp.route('/health')
def health_check():
    """Health check endpoint for container orchestration"""
    return jsonify({'status': 'healthy'}), 200


def _get_openai_client():
    """
    Lazily create an OpenAI client for Whisper STT.
    Expects OPENAI_API_KEY in environment.
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY environment variable is not set")
    return OpenAI(api_key=api_key)

@bp.route('/')
@login_required
def index():
    """Render the main chat interface"""
    try:
        # Check if user has submitted survey
        survey_model = SurveyModel(session['user_id'])
        has_submitted_survey = survey_model.has_submitted_survey()
        
        # Get user subscription tier
        from app.utils.db import get_db
        from app.models.database_models import User as DBUser
        db = get_db()
        user = db.query(DBUser).filter(DBUser.id == session['user_id']).first()
        subscription_tier = user.subscription_tier if user and user.subscription_tier else 'free'
        
        return render_template('chat.html', 
                             has_submitted_survey=has_submitted_survey,
                             subscription_tier=subscription_tier)
    except Exception as e:
        logger.error(f"Error in index route: {str(e)}")
        return render_template('chat.html', has_submitted_survey=False, subscription_tier='free')

# Add these routes to chat.py

from app.services import PromptService

@bp.route('/get_prompt')
def get_prompt():
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401
        
    try:
        prompt_service = PromptService(session['user_id'])
        current_prompt = prompt_service.get_prompt()
        return jsonify({'prompt': current_prompt})
    except Exception as e:
        logger.error(f"Error retrieving prompt: {str(e)}")
        return jsonify({'error': 'Failed to retrieve prompt'}), 500

@bp.route('/update_prompt', methods=['POST'])
def update_prompt():
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401
        
    try:
        data = request.json
        new_prompt = data.get('prompt')
   
        
        if not new_prompt:
            return jsonify({'error': 'Prompt is required'}), 400
            
        prompt_service = PromptService(session['user_id'])
        prompt_service.update_prompt(new_prompt)
        
        return jsonify({
            'success': True,
            'message': 'Prompt updated successfully'
        })
        
    except Exception as e:
        logger.error(f"Error updating prompt: {str(e)}")
        return jsonify({'error': 'Failed to update prompt'}), 500

@bp.route('/chat', methods=['POST'])
@login_required
def chat():
    """Handle chat messages and generate responses"""
    try:
        data = request.json
        user_input = data.get('input', '').strip()
        conversation_id = data.get('conversation_id')

        if not user_input:
            return jsonify({'error': 'Empty message'}), 400

        # Initialize services
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        prompt_service = PromptService(session['user_id'])

        # Get system prompt
        system_prompt = prompt_service.get_prompt()

        # Process message and generate response
        try:
            result = chat_service.process_message(
                message=user_input,
                conversation_id=conversation_id
            )
            return jsonify(result)
        except Exception as e:
            logger.error(f"Error generating response: {str(e)}")
            return jsonify({
                'error': """1:Your free key has expired,please login after 24 hours
                            2:Create another gmail account and login
                            3:Login to paid service-avalible in oct 2025"""
            }), 500

    except Exception as e:
        logger.error(f"Chat error: {str(e)}")
        return jsonify({'error': 'An error occurred'}), 500

@bp.route('/create_conversation', methods=['POST'])
@login_required
def create_conversation():
    """Create a new conversation"""
    try:
        data = request.json
        title = data.get('title', 'New Conversation')
        
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        conversation_id = chat_service.create_conversation(title)
        
        return jsonify({
            'conversation_id': conversation_id,
            'title': title
        })
    except Exception as e:
        logger.error(f"Error creating conversation: {str(e)}")
        return jsonify({'error': 'Failed to create conversation'}), 500

@bp.route('/get_conversations')
@login_required
def get_conversations():
    """Get user's recent conversations"""
    try:
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        conversations = chat_service.get_recent_conversations()
        return jsonify({'conversations': conversations})  # <-- wrap in dict for frontend
    except Exception as e:
        logger.error(f"Error retrieving conversations: {str(e)}")
        return jsonify({'error': 'Failed to retrieve conversations'}), 500

@bp.route('/get_messages/<int:conversation_id>')
@login_required
def get_messages(conversation_id):
    """Get messages for a specific conversation"""
    try:
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        messages = chat_service.get_conversation_messages(conversation_id)
        return jsonify({'messages': messages})  # <-- wrap in dict for frontend
    except Exception as e:
        logger.error(f"Error retrieving messages: {str(e)}")
        return jsonify({'error': 'Failed to retrieve messages'}), 500

@bp.route('/get_conversation/<int:conversation_id>')
@login_required
def get_conversation(conversation_id):
    """Get conversation details including title"""
    try:
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        conversation = chat_service.get_conversation_details(conversation_id)
        if conversation:
            return jsonify(conversation)
        else:
            return jsonify({'error': 'Conversation not found'}), 404
    except Exception as e:
        logger.error(f"Error retrieving conversation: {str(e)}")
        return jsonify({'error': 'Failed to retrieve conversation'}), 500

@bp.route('/delete_conversation/<int:conversation_id>', methods=['DELETE'])
@login_required
def delete_conversation(conversation_id):
    """Delete a conversation"""
    try:
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        chat_service.delete_conversation(conversation_id)
        return jsonify({'message': 'Conversation deleted successfully'})
    except Exception as e:
        logger.error(f"Error deleting conversation: {str(e)}")
        return jsonify({'error': 'Failed to delete conversation'}), 500

@bp.route('/delete_all_conversations', methods=['DELETE'])
@login_required
def delete_all_conversations():
    """Delete all conversations for the current user"""
    try:
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        chat_service.reset_all_conversations()
        return jsonify({'message': 'All conversations deleted successfully'})
    except Exception as e:
        logger.error(f"Error deleting all conversations: {str(e)}")
        return jsonify({'error': 'Failed to delete conversations'}), 500

@bp.route('/update_conversation_title/<int:conversation_id>', methods=['PUT'])
@login_required
def update_conversation_title(conversation_id):
    """Update the title of a conversation"""
    try:
        data = request.json
        new_title = data.get('title', '').strip()
        
        if not new_title:
            return jsonify({'error': 'Title cannot be empty'}), 400
        
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        success = chat_service.update_conversation_title(conversation_id, new_title)
        
        if success:
            return jsonify({'message': 'Title updated successfully', 'title': new_title})
        else:
            return jsonify({'error': 'Conversation not found or access denied'}), 404
            
    except Exception as e:
        logger.error(f"Error updating conversation title: {str(e)}")
        return jsonify({'error': 'Failed to update title'}), 500

@bp.route('/download_chat/<int:conversation_id>')
@login_required
def download_chat(conversation_id):
    """Download a chat conversation as a Word document"""
    try:
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        messages = chat_service.get_conversation_messages(conversation_id)
        
        # Create a new Word document
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('Chat Conversation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add messages
        for msg in messages:
            role = "Mr. Potter" if msg['role'] == 'bot' else "User"
            p = doc.add_paragraph()
            p.add_run(f"{role}: ").bold = True
            p.add_run(msg['message'])
            p.paragraph_format.space_after = Pt(12)
        
        # Save the document to a BytesIO object
        from io import BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Create response with appropriate headers
        from flask import make_response
        response = make_response(doc_io.getvalue())
        response.headers["Content-Disposition"] = f"attachment; filename=chat_conversation_{conversation_id}.docx"
        response.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        return response
    except Exception as e:
        logger.error(f"Error downloading chat: {str(e)}")
        return jsonify({'error': 'Failed to download chat'}), 500

@bp.route('/get_token_usage')
@login_required
def get_token_usage():
    """Get current token usage information"""
    try:
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        token_usage = chat_service.get_token_usage()
        return jsonify(token_usage)
    except Exception as e:
        logger.error(f"Error getting token usage: {str(e)}")
        return jsonify({
            'daily_limit': '100,000',
            'used_tokens': '0',
            'requested_tokens': '0',
            'wait_time': None
        }), 500
    



# TOKEN USAGE ROUTE 

@bp.route('/token_status', methods=['GET'])
@login_required
def get_token_status():
    try:
        user_id = session['user_id']
        current_api_key = session.get('groq_api_key')
        current_time = time.time()
        
        # Check cache first
        cache_key = f"{user_id}_{current_api_key}"
        if cache_key in _token_status_cache:
            cached_time, cached_data = _token_status_cache[cache_key]
            if current_time - cached_time < CACHE_TTL:
                # Return cached data
                return jsonify(cached_data)
        
        # Cache miss or expired - fetch fresh data
        # Try to get token usage directly from database first (faster)
        try:
            from app.utils.db import get_token_usage
            db_usage = get_token_usage(user_id)
            daily_limit = 100000  # Default limit
            
            used = db_usage['today']['tokens_used']
            remaining = max(0, daily_limit - used)
            
            response_data = {
                'used': used,
                'remaining': remaining,
                'limit': daily_limit,
                'reset_in': 0,  # Will be calculated by frontend if needed
                'api_key_changed': False
            }
            
            # Cache the response
            _token_status_cache[cache_key] = (current_time, response_data)
            
            # Clean old cache entries (keep only last 100)
            if len(_token_status_cache) > 100:
                # Remove oldest entries
                sorted_cache = sorted(_token_status_cache.items(), key=lambda x: x[1][0])
                for key, _ in sorted_cache[:-100]:
                    del _token_status_cache[key]
            
            return jsonify(response_data)
            
        except Exception as db_error:
            logger.warning(f"Direct DB query failed, falling back to ChatService: {str(db_error)}")
            # Fallback to ChatService if direct DB query fails
            chat_service = ChatService(user_id, current_api_key)
            
            # Verify the API key matches what's being used
            if chat_service.chat_model.api_key != current_api_key:
                chat_service.chat_model.api_key = current_api_key  # This will trigger reset
                
            token_status = chat_service.chat_model.get_token_status()
            
            response_data = {
                'used': token_status['used'],
                'remaining': token_status['remaining'],
                'limit': token_status['limit'],
                'reset_in': token_status['reset_in'],
                'api_key_changed': False  # Can be used by frontend if needed
            }
            
            # Cache the response
            _token_status_cache[cache_key] = (current_time, response_data)
            
            # Clean old cache entries (keep only last 100)
            if len(_token_status_cache) > 100:
                # Remove oldest entries
                sorted_cache = sorted(_token_status_cache.items(), key=lambda x: x[1][0])
                for key, _ in sorted_cache[:-100]:
                    del _token_status_cache[key]
            
            return jsonify(response_data)
            
    except Exception as e:
        logger.error(f"Error getting token status: {str(e)}", exc_info=True)
        # Return cached data if available, even if expired
        cache_key = f"{session.get('user_id')}_{session.get('groq_api_key')}"
        if cache_key in _token_status_cache:
            _, cached_data = _token_status_cache[cache_key]
            return jsonify(cached_data)
        # Return error response
        return jsonify({
            'error': 'Failed to get token status',
            'used': 0,
            'remaining': 0,
            'limit': 100000,
            'reset_in': 0
        }), 500

@bp.route('/user_info')
def get_user_info():
    """Get current user information including role"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401
    
    try:
        from app.models import UserModel
        user_model = UserModel(session['user_id'])
        user_info = UserModel.get_user_by_id(session['user_id'])
        
        if not user_info:
            return jsonify({'error': 'User not found'}), 404
        
        return jsonify({
            'success': True,
            'user': {
                'id': user_info['id'],
                'username': user_info['username'],
                'role': user_info.get('role', 'student'),
                'class_standard': user_info['class_standard'],
                'medium': user_info['medium']
            }
        })
    except Exception as e:
        logger.error(f"Error getting user info: {str(e)}")
        return jsonify({'error': 'Failed to get user info'}), 500


@bp.route('/api/stt', methods=['POST'])
@login_required
def speech_to_text():
    """
    Convert uploaded speech audio to text using OpenAI Whisper.

    Expects multipart/form-data with field "audio".
    Returns JSON: {"text": "..."} on success.
    """
    try:
        if 'audio' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400

        audio_file = request.files['audio']
        if audio_file.filename == '':
            return jsonify({'error': 'Empty audio filename'}), 400

        # OpenAI client for Whisper
        client = _get_openai_client()

        # Whisper works best with binary file-like objects
        # We read into memory here as recordings are short (voice messages)
        from tempfile import NamedTemporaryFile

        with NamedTemporaryFile(delete=False, suffix=".webm") as tmp:
            audio_file.save(tmp.name)
            tmp_path = tmp.name

        with open(tmp_path, "rb") as f:
            transcription = client.audio.transcriptions.create(
                model="whisper-1",
                file=f,
                response_format="json"
            )

        text = getattr(transcription, "text", None) or transcription.get("text")  # handle both object/dict
        if not text:
            return jsonify({'error': 'Transcription failed'}), 500

        return jsonify({'text': text})

    except Exception as e:
        logger.error(f"Error in speech_to_text: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to transcribe audio'}), 500

@bp.route('/chatbot', methods=['GET'])
@teacher_required
def chatbot():
    """Render the chatbot interface"""
    try:
        # Get user's lessons for selection
        user_id = session.get('user_id')
        lessons = LessonModel.get_lessons_by_teacher(user_id)
        
        return render_template('chatbot.html', lessons=lessons)
    except Exception as e:
        logger.error(f"Error rendering chatbot: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to render chatbot: {str(e)}'}), 500


@bp.route('/chatbot_update', methods=['GET'])
@login_required
def chatbot_update():
    """Render the PDF chat interface"""
    try:
        return render_template('chatbot_update.html')
    except Exception as e:
        logger.error(f"Error rendering chatbot_update page: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to render page: {str(e)}'}), 500




import re
from io import BytesIO
from flask import request, jsonify, send_file
from gtts import gTTS
from langdetect import detect

def clean_text_for_tts(text: str) -> str:
    """
    Remove markdown and unwanted symbols so TTS only speaks readable text.
    Keeps letters (including Urdu/Arabic), numbers, punctuation, and spaces.
    """
    # Remove markdown & special symbols
    text = re.sub(r"[#*_~`>|=\[\]{}()^]", "", text)

    # Remove multiple spaces
    text = re.sub(r"\s+", " ", text)

    return text.strip()

@bp.route('/api/tts', methods=['POST'])
@login_required
def text_to_speech():
    """
    Convert text to speech using gTTS.
    Cleans symbols before sending text to TTS.
    """
    try:
        data = request.get_json() or {}
        text = (data.get('text') or '').strip()

        if not text:
            return jsonify({'error': 'Text is required'}), 400

        # âœ… Clean text before TTS
        text = clean_text_for_tts(text)

        # Detect language; fallback to English
        try:
            lang = detect(text)
        except Exception as e:
            logger.warning(f"Language detection failed, defaulting to 'en': {str(e)}")
            lang = 'en'

        # Generate speech
        tts = gTTS(text=text, lang=lang)
        audio_fp = BytesIO()
        tts.write_to_fp(audio_fp)
        audio_fp.seek(0)

        return send_file(
            audio_fp,
            mimetype='audio/mpeg',
            as_attachment=False,
            download_name='tts.mp3'
        )

    except Exception as e:
        logger.error(f"Error in text_to_speech: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to generate speech'}), 500



# def text_to_speech():
#     """
#     Convert text to speech using gTTS and return an audio stream.
#     This replaces browser-based SpeechSynthesis so it also works for Urdu and other languages.
#     """
#     try:
#         data = request.get_json() or {}
#         text = (data.get('text') or '').strip()

#         if not text:
#             return jsonify({'error': 'Text is required'}), 400

#         # Detect language; fallback to English on failure
#         try:
#             lang = detect(text)
#         except Exception as e:
#             logger.warning(f"Language detection failed, defaulting to 'en': {str(e)}")
#             lang = 'en'

#         # Generate speech with gTTS
#         tts = gTTS(text=text, lang=lang)
#         audio_fp = BytesIO()
#         tts.write_to_fp(audio_fp)
#         audio_fp.seek(0)

#         # Return audio as an MP3 stream
#         return send_file(
#             audio_fp,
#             mimetype='audio/mpeg',
#             as_attachment=False,
#             download_name='tts.mp3'
#         )
#     except Exception as e:
#         logger.error(f"Error in text_to_speech: {str(e)}", exc_info=True)
#         return jsonify({'error': 'Failed to generate speech'}), 500