from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

def create_title(document, text):
    """Add a title to the document"""
    title = document.add_heading(text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

def add_section_heading(document, text, level=1):
    """Add a section heading"""
    heading = document.add_heading(text, level)
    if level == 1:
        heading.style.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue

def add_code_block(document, code, language=""):
    """Add a formatted code block"""
    p = document.add_paragraph()
    if language:
        p.add_run(f"{language}:\n").bold = True
    code_run = p.add_run(code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    document.add_paragraph()

def generate_implementation_notes():
    """Generate implementation notes document"""
    doc = Document()
    
    # Title
    create_title(doc, "Implementation Notes")
    
    # Add date
    date_paragraph = doc.add_paragraph()
    date_paragraph.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}").bold = True
    doc.add_paragraph()

    # 1. Email Verification System
    add_section_heading(doc, "1. Email Verification System Implementation")

    # A. register_email.html
    add_section_heading(doc, "A. register_email.html Implementation", 2)
    add_code_block(doc, """- Complete HTML structure with Bootstrap integration
- Form components:
  • Email input field with validation
  • Submit button for verification email
  • Error message display section
  • Login link for existing users
- Styling:
  • Custom error message styling
  • Responsive form container
  • Bootstrap classes for layout""", "HTML")

    # B. email_sent.html
    add_section_heading(doc, "B. email_sent.html Implementation", 2)
    add_code_block(doc, """- Confirmation page structure:
  • Success icon using Font Awesome
  • Verification email sent message
  • Email check prompt
  • Spam folder check alert
- Features:
  • Resend verification option
  • Link back to registration
  • Responsive design with Bootstrap""", "HTML")

    # C. register.html
    add_section_heading(doc, "C. register.html Updates", 2)
    add_code_block(doc, """- Title update to "Register - Complete Profile"
- New components:
  • Hidden email input for verified address
  • Verified email display section
  • Updated form submission flow
- Visual feedback:
  • Success indicators for verified email
  • Error message handling
  • Updated button text""", "HTML")

    # 2. Survey System Implementation
    add_section_heading(doc, "2. Survey System Implementation")

    # A. Database Schema
    add_section_heading(doc, "A. Database Schema", 2)
    add_code_block(doc, """CREATE TABLE user_surveys (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    user_id INTEGER NOT NULL,
    experience_rating TEXT NOT NULL,
    is_helpful BOOLEAN NOT NULL,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (user_id) REFERENCES users (id)
);""", "SQL")

    # B. Survey Modal
    add_section_heading(doc, "B. Survey Modal Implementation", 2)
    add_code_block(doc, """<div class="modal fade" id="surveyModal">
  - Experience Rating Section:
    • Positive/Neutral/Negative options
    • Radio button implementation
    • Required field validation
  - Helpfulness Section:
    • Yes/No options
    • Required field validation
  - Submit button with event handler""", "HTML")

    # C. JavaScript Implementation
    add_section_heading(doc, "C. Survey JavaScript Implementation", 2)
    add_code_block(doc, """- State Management:
  • hasSubmittedSurvey flag
  • isLoggingOut flag

- Core Functions:
  • showSurveyModal()
  • submitSurvey()
  • Event handlers for:
    - Logout button
    - Page unload
    - Form submission

- API Integration:
  • POST request to /submit_survey
  • FormData handling
  • Response processing
  • Error handling""", "JavaScript")

    # 3. PDF Document Handling
    add_section_heading(doc, "3. PDF Document Handling Improvements")

    # A. ChatService Enhancement
    add_section_heading(doc, "A. Enhanced ChatService Class", 2)
    add_code_block(doc, """class ChatService:
    def get_document_context(self, message: str) -> str:
        # Page Detection
        page_match = re.search(r'page\s*(\d+)', message.lower())
        page_number = int(page_match.group(1)) if page_match else None
        
        # Context Retrieval
        if self.vector_store._vectorstore:
            relevant_docs = self.vector_store.search_similar(message, k=3)
            
            # Document Processing
            for doc in relevant_docs:
                page = doc.metadata.get('page', 1)
                if page_number and page != page_number:
                    continue
                context_parts.append(f"Page {page}:\\n{doc.page_content.strip()}")""", "Python")

    # B. System Prompts
    add_section_heading(doc, "B. System Prompt Enhancements", 2)
    add_code_block(doc, """base_prompt = (
    "You are a helpful assistant for the current conversation only. "
    "When referring to document content, always mention the page number and filename. "
    "If asked about specific pages, focus on content from those pages. "
    "If no relevant content is found for a specific page, clearly state that. "
    "Format your responses in a clear, organized manner. "
    "If any one asks about the data say the data is crawled from the internet "
    "if user not ask about the data then no need to say the data is crawled from the internet"
)""", "Python")

    # 4. Integration Testing
    add_section_heading(doc, "4. Integration Testing Points")
    add_code_block(doc, """A. Email Verification Flow:
   - Email submission → Verification email → Registration completion

B. Survey Trigger Points:
   - Logout button click
   - Browser/tab close
   - Form submission

C. PDF Document Queries:
   - General content search
   - Page-specific queries
   - Error cases""")

    # 5. Security Considerations
    add_section_heading(doc, "5. Security Considerations")
    add_code_block(doc, """A. Email Verification:
   - Token-based verification
   - Expiration handling
   - XSS prevention in templates

B. Survey Data:
   - User authentication required
   - Input validation
   - CSRF protection

C. Document Handling:
   - Access control
   - Input sanitization
   - Error message security""")

    # Save the document
    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs')
    os.makedirs(output_dir, exist_ok=True)
    doc.save(os.path.join(output_dir, f'implementation_notes_{datetime.now().strftime("%Y%m%d")}.docx'))

if __name__ == "__main__":
    generate_implementation_notes() 